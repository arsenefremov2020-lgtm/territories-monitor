import os
import hashlib
import requests
import pandas as pd
from sqlalchemy import create_engine, text
from bs4 import BeautifulSoup
from urllib.parse import urljoin
from docx import Document
from io import BytesIO

DATABASE_URL = os.environ["DATABASE_URL"]
PRINT_URL = "https://zakon.rada.gov.ua/laws/show/z0380-25/print"

engine = create_engine(DATABASE_URL)

CATEGORY_BY_TABLE = {
    1: "Території можливих бойових дій",
    2: "Території бойових дій",
    3: "Території активних бойових дій",
    4: "Тимчасово окуповані території",
    5: "Тимчасово окуповані території",
}


def get_docx_url():
    html = requests.get(PRINT_URL, timeout=30).text
    soup = BeautifulSoup(html, "html.parser")

    for a in soup.find_all("a", href=True):
        href = a["href"]
        if href.endswith(".docx"):
            return urljoin(PRINT_URL, href)

    raise Exception("DOCX link not found")


def download_docx(docx_url):
    response = requests.get(docx_url, timeout=60)
    response.raise_for_status()
    file_hash = hashlib.sha256(response.content).hexdigest()
    return response.content, file_hash


def normalize_date(value):
    value = str(value).strip()

    if value == "" or value.lower() in ["none", "nan"]:
        return None

    parts = value.split(".")
    if len(parts) == 3:
        return f"{parts[2]}-{parts[1]}-{parts[0]}"

    return None


def is_territory_code(value):
    value = str(value).strip()
    return value.startswith("UA") and len(value) >= 9


def is_oblast_row(value):
    value = str(value).strip()
    return "ОБЛАСТЬ" in value.upper() or "АВТОНОМНА РЕСПУБЛІКА КРИМ" in value.upper()


def is_rayon_row(value):
    value = str(value).strip()
    return "район" in value.lower()


def clean_repeated_header(value):
    import re

    value = str(value).replace("\n", " ").strip()
    value = re.sub(r"\s+", " ", value)

    parts = [p.strip() for p in value.split("  ") if p.strip()]
    if parts:
        value = parts[0]

    value = re.sub(r"^\d+(\.\d+)*\.\s*", "", value)
    value = re.sub(r"^\d+(\.\d+)*\s*", "", value)

    return value.strip()


def extract_rows(docx_content):
    document = Document(BytesIO(docx_content))
    if len(document.tables) < 5:
        raise Exception(f"Unexpected DOCX structure: expected at least 5 tables, found {len(document.tables)}")
    rows = []

    for table_index, table in enumerate(document.tables, start=1):
        category = CATEGORY_BY_TABLE.get(table_index, f"Таблиця {table_index}")

        current_oblast = None
        current_rayon = None

        for row in table.rows:
            cells = [cell.text.replace("\n", " ").strip() for cell in row.cells]

            if len(cells) < 4:
                continue
            if len(cells) != 4:
    print("Warning: unexpected column count:", len(cells), cells)

            first_cell = clean_repeated_header(cells[0])
            second_cell = clean_repeated_header(cells[1])

            if is_oblast_row(first_cell):
                current_oblast = first_cell.title()
                current_rayon = None
                continue

            if is_rayon_row(first_cell) and not is_territory_code(first_cell):
                current_rayon = first_cell
                continue

            full_code = cells[0].strip()
            territory_name = cells[1].strip()
            status_from = cells[2].strip()
            status_to = cells[3].strip()

            if not is_territory_code(full_code):
                continue

            rows.append({
                "full_code": full_code,
                "hromada_code_7": full_code.replace("UA", "")[:7],
                "territory_name": territory_name,
                "oblast": current_oblast,
                "rayon": current_rayon,
                "category": category,
                "status_from": normalize_date(status_from),
                "status_to": normalize_date(status_to),
            })

    return pd.DataFrame(rows)


def save_document_record(docx_url, file_hash):
    with engine.begin() as conn:
        existing = conn.execute(
            text("select id from documents where file_hash = :file_hash"),
            {"file_hash": file_hash},
        ).fetchone()

        if existing:
            return existing[0], False

        result = conn.execute(
            text("""
                insert into documents
                (source_name, source_url, document_number, document_date, file_hash)
                values
                (:source_name, :source_url, :document_number, :document_date, :file_hash)
                returning id
            """),
            {
                "source_name": "Верховна Рада України / zakon.rada.gov.ua",
                "source_url": docx_url,
                "document_number": "376",
                "document_date": "2025-02-28",
                "file_hash": file_hash,
            },
        )

        return result.fetchone()[0], True


def replace_current_data(df, document_id):
    df["source_document_id"] = document_id

    with engine.begin() as conn:
        conn.execute(text("delete from territory_status_history"))

    df.to_sql(
        "territory_status_history",
        engine,
        if_exists="append",
        index=False
    )


if __name__ == "__main__":
    docx_url = get_docx_url()
    print("DOCX URL:", docx_url)

    docx_content, file_hash = download_docx(docx_url)
    print("File hash:", file_hash)

    document_id, is_new = save_document_record(docx_url, file_hash)
    print("Document ID:", document_id)
    print("Is new document:", is_new)

    df = extract_rows(docx_content)
    print("Rows extracted:", len(df))
    print(df.head(20))

    replace_current_data(df, document_id)
    print("Database updated successfully")
